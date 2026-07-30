"""Sinh bản `.docx` của báo cáo A2 từ `docs/bao-cao-a2.md`.

Bản Markdown là nguồn sự thật; file này chỉ đổi định dạng, không đổi nội dung.
Sửa nội dung thì sửa Markdown rồi chạy lại, đừng sửa trong `.docx`.

    python tools/bao-cao-sang-docx.py

Yêu cầu: python-docx. Bản `.docx` ra ở gốc kho mã và cố ý không được commit,
vì nó là sản phẩm dẫn xuất; xem `.gitignore`.

Script chỉ hiểu tập cú pháp Markdown mà báo cáo đang dùng: tiêu đề, đoạn văn,
bảng, khối mã, khối trích dẫn, và ba dạng đánh dấu trong dòng là đậm, nghiêng,
mã. Danh sách gạch đầu dòng và danh sách đánh số **không** nằm trong tập đó, và
chúng hỏng âm thầm chứ không báo lỗi: các dòng của một danh sách sẽ bị gộp thành
một đoạn văn liền. Thêm một danh sách vào báo cáo thì phải dạy script trước.

Số trang thật chỉ đếm được bằng cách mở trong Word. Con số script in ra là số
ước, tính từ một mô hình đếm dòng, nên nó dùng để biết cần siết hay nới chứ
không dùng để nghiệm thu.

Nếu số trang lệch khỏi khoảng 12 tới 15 thì đổi **một** giá trị trong khối NÚM
ĐIỀU CHỈNH dưới đây rồi chạy lại; đừng cắt chữ trong báo cáo để chạy theo số
trang. Nới ra: GIAN_DONG 1,0 -> 1,15. Siết vào: CO_CHU 12 -> 11,5 hoặc LE 2,5 -> 2,2.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# --- NÚM ĐIỀU CHỈNH ---------------------------------------------------------
PHONG_CHU = "Times New Roman"
CO_CHU = 12  # pt
GIAN_DONG = 1.0  # 1,0 là dòng đơn; 1,15 nới ra khoảng một trang trên mười
LE = 2.5  # cm, cả bốn phía
CO_CHU_BANG = 11
CO_CHU_MA = 10
# ---------------------------------------------------------------------------

GOC = Path(__file__).resolve().parent.parent
NGUON = GOC / "docs" / "bao-cao-a2.md"
DICH = GOC / "bao-cao-a2.docx"

# Ba dạng đánh dấu trong dòng, gộp thành một biểu thức để quét một lượt.
TRONG_DONG = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def dat_kieu(tai_lieu):
    thuong = tai_lieu.styles["Normal"]
    thuong.font.name = PHONG_CHU
    thuong.font.size = Pt(CO_CHU)
    # Word chọn font cho chữ có dấu theo thuộc tính eastAsia, không theo ascii,
    # nên thiếu dòng này thì tiếng Việt có dấu bị đổi sang font khác.
    thuong.element.rPr.rFonts.set(qn("w:eastAsia"), PHONG_CHU)
    doan = thuong.paragraph_format
    doan.line_spacing = GIAN_DONG
    doan.space_after = Pt(4)
    doan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for ten, co, truoc in (("Title", 15, 0), ("Heading 1", 14, 12), ("Heading 2", 12, 8)):
        kieu = tai_lieu.styles[ten]
        kieu.font.name = PHONG_CHU
        kieu.font.size = Pt(co)
        kieu.font.bold = True
        kieu.font.color.rgb = None
        kieu.element.rPr.rFonts.set(qn("w:eastAsia"), PHONG_CHU)
        kieu.paragraph_format.space_before = Pt(truoc)
        kieu.paragraph_format.space_after = Pt(4)
        kieu.paragraph_format.line_spacing = GIAN_DONG
        kieu.paragraph_format.keep_with_next = True

    for muc in tai_lieu.sections:
        muc.top_margin = muc.bottom_margin = Cm(LE)
        muc.left_margin = muc.right_margin = Cm(LE)


def so_trang_o_chan(tai_lieu):
    """Chèn trường số trang vào chân trang, để đếm trang bằng mắt được."""
    doan = tai_lieu.sections[0].footer.paragraphs[0]
    doan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    o = doan.add_run()
    for thuoc_tinh, gia_tri in (("w:fldCharType", "begin"), (None, "PAGE"), ("w:fldCharType", "end")):
        if thuoc_tinh is None:
            phan_tu = o._element.makeelement(qn("w:instrText"), {})
            phan_tu.text = f" {gia_tri} "
        else:
            phan_tu = o._element.makeelement(qn("w:fldChar"), {qn(thuoc_tinh): gia_tri})
        o._element.append(phan_tu)
    o.font.size = Pt(CO_CHU_BANG)


def viet_dong(doan, van_ban, dam=False, nghieng=False):
    """Rải một dòng Markdown vào một đoạn, giữ ba dạng đánh dấu trong dòng.

    Gọi lại chính nó cho phần nằm trong một cặp đánh dấu, vì báo cáo có chỗ
    lồng nhau như **`docker build` không tái lập được.**; bóc một lớp rồi dừng
    thì lớp bên trong lọt ra ngoài ở dạng ký tự thô.
    """
    for phan in TRONG_DONG.split(van_ban):
        if not phan:
            continue
        if phan.startswith("**") and phan.endswith("**"):
            viet_dong(doan, phan[2:-2], dam=True, nghieng=nghieng)
        elif phan.startswith("`") and phan.endswith("`"):
            o = doan.add_run(phan[1:-1])
            o.font.name = "Consolas"
            o.font.size = Pt(CO_CHU_BANG)
            o.bold, o.italic = dam, nghieng
        elif phan.startswith("*") and phan.endswith("*"):
            viet_dong(doan, phan[1:-1], dam=dam, nghieng=True)
        else:
            o = doan.add_run(phan)
            o.bold, o.italic = dam, nghieng


def mo_khoi_khac(dong):
    """Dòng này có mở một khối không phải văn xuôi không.

    Một dấu ` ở đầu dòng là mã trong dòng chứ không phải khối mã, nên chỉ ba
    dấu backtick liền mới tính; nhận nhầm chỗ này thì đoạn văn bắt đầu bằng
    một tên file sẽ không cắt được.
    """
    return dong.startswith(("#", ">", "|", "```"))


def cat_khoi(dong_md):
    """Cắt Markdown thành các khối, mỗi khối là (loại, các dòng)."""
    khoi, i = [], 0
    while i < len(dong_md):
        dong = dong_md[i]
        if not dong.strip():
            i += 1
        elif dong.startswith("```"):
            j = i + 1
            while j < len(dong_md) and not dong_md[j].startswith("```"):
                j += 1
            khoi.append(("ma", dong_md[i + 1 : j]))
            i = j + 1
        elif dong.startswith("#"):
            khoi.append(("tieu-de", [dong]))
            i += 1
        elif dong.lstrip().startswith("|"):
            j = i
            while j < len(dong_md) and dong_md[j].lstrip().startswith("|"):
                j += 1
            khoi.append(("bang", dong_md[i:j]))
            i = j
        elif dong.startswith(">"):
            j = i
            while j < len(dong_md) and dong_md[j].startswith(">"):
                j += 1
            khoi.append(("trich", dong_md[i:j]))
            i = j
        else:
            j = i + 1
            while j < len(dong_md) and dong_md[j].strip() and not mo_khoi_khac(dong_md[j]):
                j += 1
            khoi.append(("doan", dong_md[i:j]))
            i = j
    return khoi


def dung_bang(tai_lieu, dong_bang):
    o_dong = [[o.strip() for o in d.strip().strip("|").split("|")] for d in dong_bang]
    # Dòng thứ hai của một bảng GFM là dòng gạch ngăn, không phải dữ liệu.
    dau, than = o_dong[0], o_dong[2:]
    bang = tai_lieu.add_table(rows=1 + len(than), cols=len(dau))
    bang.style = "Table Grid"
    bang.alignment = WD_TABLE_ALIGNMENT.CENTER
    for hang_o, hang in zip([dau] + than, bang.rows):
        for van_ban, o in zip(hang_o, hang.cells):
            doan = o.paragraphs[0]
            doan.paragraph_format.space_after = Pt(2)
            doan.paragraph_format.line_spacing = 1.0
            doan.alignment = WD_ALIGN_PARAGRAPH.LEFT
            viet_dong(doan, van_ban)
            for chay in doan.runs:
                chay.font.size = Pt(CO_CHU_BANG)
                if hang_o is dau:
                    chay.bold = True
    return bang


def sinh():
    # Console mặc định của Windows là cp1252 và nó không mã hoá được tiếng Việt
    # có dấu, nên thiếu dòng này thì script chết ở lệnh print cuối cùng, sau khi
    # đã ghi xong file, và trông như một lần chạy hỏng.
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    if not NGUON.exists():
        sys.exit(f"Không thấy {NGUON}")
    dong_md = NGUON.read_text(encoding="utf-8").split("\n")

    tai_lieu = Document()
    dat_kieu(tai_lieu)
    so_trang_o_chan(tai_lieu)

    for loai, dong in cat_khoi(dong_md):
        if loai == "tieu-de":
            muc = len(dong[0]) - len(dong[0].lstrip("#"))
            kieu = {1: "Title", 2: "Heading 1"}.get(muc, "Heading 2")
            doan = tai_lieu.add_paragraph(style=kieu)
            if muc == 1:
                doan.alignment = WD_ALIGN_PARAGRAPH.CENTER
            viet_dong(doan, dong[0].lstrip("#").strip())
        elif loai == "bang":
            dung_bang(tai_lieu, dong)
            tai_lieu.add_paragraph().paragraph_format.space_after = Pt(2)
        elif loai == "ma":
            doan = tai_lieu.add_paragraph()
            doan.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doan.paragraph_format.line_spacing = 1.0
            doan.paragraph_format.left_indent = Cm(0.5)
            o = doan.add_run("\n".join(dong))
            o.font.name = "Consolas"
            o.font.size = Pt(CO_CHU_MA)
        elif loai == "trich":
            doan = tai_lieu.add_paragraph()
            doan.paragraph_format.left_indent = Cm(1)
            viet_dong(doan, " ".join(d.lstrip("> ").rstrip() for d in dong))
        else:
            van_ban = " ".join(d.strip() for d in dong)
            doan = tai_lieu.add_paragraph()
            viet_dong(doan, van_ban)
            # Một đoạn chỉ gồm đúng một cụm in đậm là nhãn của khối ngay dưới
            # nó, ví dụ bốn nhãn của mỗi mục con trong mục 3. Không giữ lại thì
            # có lúc nhãn nằm một mình ở cuối trang còn khối thì sang trang sau.
            if van_ban.startswith("**") and van_ban.endswith("**") and van_ban.count("**") == 2:
                doan.paragraph_format.keep_with_next = True

    tai_lieu.save(DICH)

    # Ước số trang bằng cách đếm dòng, không phải bằng cách phân trang thật.
    van_ban = "\n".join(dong_md)
    tu = len(re.findall(r"\S+", van_ban))
    rong_dong = (21.0 - 2 * LE) * 10 / (CO_CHU * 0.176)  # số ký tự lọt một dòng
    cao_trang = (29.7 - 2 * LE) * 10 / (CO_CHU * GIAN_DONG * 0.407)  # số dòng lọt một trang
    so_dong = sum(max(1, round(len(d) / rong_dong + 0.5)) for d in dong_md if d.strip())
    print(f"Đã ghi {DICH}")
    print(f"{tu} từ, ước {so_dong / cao_trang:.1f} trang")
    print("Số trang thật: mở trong Word và đọc chân trang cuối.")


if __name__ == "__main__":
    sinh()
